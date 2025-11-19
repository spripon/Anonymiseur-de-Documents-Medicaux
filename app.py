import streamlit as st
import re
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from datetime import datetime
from PIL import Image, ImageDraw
import pytesseract
import numpy as np
import cv2

# ------------------------------------------------------------
# Configuration de la page
# ------------------------------------------------------------
st.set_page_config(
    page_title="Anonymiseur de Documents Médicaux",
    page_icon="🏥",
    layout="wide",
)

st.title("🏥 Anonymiseur de Documents Médicaux")
st.markdown("---")

# ------------------------------------------------------------
# Patterns de détection
# ------------------------------------------------------------
PATTERNS = {
    "dates": r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    "numeros_longs": r"\b\d{6,}\b",
    "noms_propres": r"\b[A-ZÉÈÊËÀÂÄÔÖÛÜÇ][a-zéèêëàâäôöûüç]+(?:\s+[A-ZÉÈÊËÀÂÄÔÖÛÜÇ][a-zéèêëàâäôöûüç]+)*\b",
    "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
    "telephone": r"\b(?:\+33|0)[1-9](?:[\s.-]?\d{2}){4}\b",
    "numero_secu": r"\b[12]\s?\d{2}\s?\d{2}\s?\d{2}\s?\d{3}\s?\d{3}\s?\d{2}\b",
}

LABELS_COMMUNS = [
    "Nom", "Prenom", "N° patient", "Numero patient", "Patient",
    "Age", "Date de naissance", "Ne(e) le",
    "Etablissement", "Hopital", "Clinique",
    "Date etude", "Date d'etude", "Date examen",
    "Effectue par", "Realise par", "Medecin", "Docteur", "Dr",
    "Adresse", "Telephone", "Tel", "Email", "N°SS", "Securite sociale",
]


# ------------------------------------------------------------
# Fonctions d'anonymisation (texte)
# ------------------------------------------------------------
def anonymize_text(text: str, labels_to_remove):
    anonymized = text
    replacements = []

    # Dates
    for match in re.finditer(PATTERNS["dates"], text):
        original = match.group()
        anonymized = anonymized.replace(original, "[DATE ANONYMISEE]")
        replacements.append(("Date", original, "[DATE ANONYMISEE]"))

    # Numéros longs
    for match in re.finditer(PATTERNS["numeros_longs"], text):
        original = match.group()
        if not re.search(r"\d{1,2}[/-]" + re.escape(original), text):
            anonymized = anonymized.replace(original, "[NUMERO ANONYMISE]")
            replacements.append(("Numero", original, "[NUMERO ANONYMISE]"))

    # Emails
    for match in re.finditer(PATTERNS["email"], text):
        original = match.group()
        anonymized = anonymized.replace(original, "[EMAIL ANONYMISE]")
        replacements.append(("Email", original, "[EMAIL ANONYMISE]"))

    # Téléphones
    for match in re.finditer(PATTERNS["telephone"], text):
        original = match.group()
        anonymized = anonymized.replace(original, "[TEL ANONYMISE]")
        replacements.append(("Telephone", original, "[TEL ANONYMISE]"))

    # Numéros de sécurité sociale
    for match in re.finditer(PATTERNS["numero_secu"], text):
        original = match.group()
        anonymized = anonymized.replace(original, "[N°SS ANONYMISE]")
        replacements.append(("N°SS", original, "[N°SS ANONYMISE]"))

    # Labels structurés (Nom: ..., Prenom: ..., etc.)
    for label in labels_to_remove:
        pattern = rf"{re.escape(label)}\s*:?\s*([^\n]+)"
        for match in re.finditer(pattern, anonymized, re.IGNORECASE):
            full_match = match.group(0)
            value = match.group(1).strip()
            if value:
                replacement = f"{label}: [ANONYMISE]"
                anonymized = anonymized.replace(full_match, replacement)
                replacements.append((label, value, "[ANONYMISE]"))

    return anonymized, replacements


def anonymize_pdf(pdf_bytes: bytes, labels_to_remove):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    all_replacements = []

    text_preserve_flag = getattr(fitz, "TEXT_PRESERVE_WHITESPACE", 0)

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()

        anonymized_text, replacements = anonymize_text(text, labels_to_remove)
        all_replacements.extend(replacements)

        # Masquage par zone autour des labels
        for label in labels_to_remove:
            areas = page.search_for(label, flags=text_preserve_flag)
            for area in areas:
                extended_area = fitz.Rect(area.x0, area.y0, area.x0 + 300, area.y1)
                page.add_redact_annot(extended_area, fill=(0, 0, 0))

        # Masquage visuel des occurrences détectées par regex
        for pattern_key in ["dates", "numeros_longs", "email", "telephone", "numero_secu"]:
            for match in re.finditer(PATTERNS[pattern_key], text):
                for area in page.search_for(match.group()):
                    page.add_redact_annot(area, fill=(0, 0, 0))

        page.apply_redactions()

    # Sauvegarde en mémoire
    output_buffer = BytesIO()
    doc.save(output_buffer)
    doc.close()
    output_buffer.seek(0)

    return output_buffer.getvalue(), all_replacements


def anonymize_docx(docx_bytes: bytes, labels_to_remove):
    doc = Document(BytesIO(docx_bytes))
    all_replacements = []

    # Paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            anonymized_text, replacements = anonymize_text(para.text, labels_to_remove)
            all_replacements.extend(replacements)
            para.text = anonymized_text

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    anonymized_text, replacements = anonymize_text(cell.text, labels_to_remove)
                    all_replacements.extend(replacements)
                    cell.text = anonymized_text

    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    return output_buffer.getvalue(), all_replacements


def anonymize_txt(txt_bytes: bytes, labels_to_remove):
    text = txt_bytes.decode("utf-8", errors="ignore")
    anonymized_text, replacements = anonymize_text(text, labels_to_remove)
    return anonymized_text.encode("utf-8"), replacements


# ------------------------------------------------------------
# Fonctions d'anonymisation (images)
# ------------------------------------------------------------
def anonymize_image(image_bytes: bytes, labels_to_remove, use_ocr: bool = True):
    # Chargement initial
    image = Image.open(BytesIO(image_bytes))

    # Normalisation du mode et du format pour pytesseract
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")

    # On force un format supporté pour éviter "Unsupported image format/type"
    image.format = "PNG"

    anonymized_image = image.copy()
    draw = ImageDraw.Draw(anonymized_image)
    all_replacements = []

    # OCR
    if use_ocr:
        try:
            ocr_data = pytesseract.image_to_data(
                anonymized_image,
                lang="fra+eng",
                output_type=pytesseract.Output.DICT
            )

            n_boxes = len(ocr_data["text"])
            for i in range(n_boxes):
                text = ocr_data["text"][i].strip()
                if not text:
                    continue

                try:
                    conf = int(ocr_data["conf"][i])
                except ValueError:
                    continue

                if conf <= 30:
                    continue

                should_anonymize = False
                replacement_type = ""

                if re.match(PATTERNS["dates"], text):
                    should_anonymize = True
                    replacement_type = "Date"
                elif re.match(PATTERNS["numeros_longs"], text):
                    should_anonymize = True
                    replacement_type = "Numero"
                elif re.match(PATTERNS["email"], text):
                    should_anonymize = True
                    replacement_type = "Email"
                elif re.match(PATTERNS["telephone"], text):
                    should_anonymize = True
                    replacement_type = "Telephone"
                else:
                    for label in labels_to_remove:
                        if label.lower() in text.lower():
                            should_anonymize = True
                            replacement_type = label
                            break

                if should_anonymize:
                    x = ocr_data["left"][i]
                    y = ocr_data["top"][i]
                    w = ocr_data["width"][i]
                    h = ocr_data["height"][i]

                    padding = 5
                    x -= padding
                    y -= padding
                    w += padding * 2
                    h += padding * 2

                    draw.rectangle([x, y, x + w, y + h], fill="black")
                    all_replacements.append((replacement_type, text, "[ANONYMISE]"))

        except Exception as e:
            st.warning(f"OCR non disponible ou erreur: {str(e)}. Anonymisation manuelle appliquee.")

    # Détection simple de zones textuelles / en-têtes via OpenCV
    try:
        img_array = np.array(anonymized_image)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

        thresh = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            11,
            2,
        )

        contours, _ = cv2.findContours(
            thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )

        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)

            if 20 < w < anonymized_image.width * 0.8 and 10 < h < 100:
                roi = gray[y:y + h, x:x + w]
                white_pixel_ratio = float(np.sum(roi > 200)) / float(w * h)

                if 0.3 < white_pixel_ratio < 0.95 and y < anonymized_image.height * 0.3:
                    draw.rectangle([x, y, x + w, y + h], fill="black")
                    all_replacements.append(
                        ("Zone detectee", f"Position ({x},{y})", "[MASQUE]")
                    )

    except Exception as e:
        st.warning(f"Detection automatique de zones limitee: {str(e)}")

    # Sauvegarde de l'image anonymisée
    output_buffer = BytesIO()
    img_format = anonymized_image.format if anonymized_image.format else "PNG"
    anonymized_image.save(output_buffer, format=img_format)
    output_buffer.seek(0)

    return output_buffer.getvalue(), all_replacements, img_format


# ------------------------------------------------------------
# Barre latérale : configuration
# ------------------------------------------------------------
st.sidebar.header("Configuration")

st.sidebar.subheader("Labels à anonymiser")
selected_labels = st.sidebar.multiselect(
    "Sélectionnez les champs à anonymiser :",
    LABELS_COMMUNS,
    default=[
        "Nom", "Prenom", "N° patient", "Age", "Date de naissance",
        "Etablissement", "Date etude", "Effectue par",
    ],
)

custom_labels = st.sidebar.text_area(
    "Labels personnalisés (un par ligne) :",
    help="Ajoutez des labels supplémentaires à anonymiser",
)

if custom_labels:
    custom_labels_list = [
        label.strip() for label in custom_labels.split("\n") if label.strip()
    ]
    selected_labels.extend(custom_labels_list)

st.sidebar.subheader("Options pour les images")
use_ocr = st.sidebar.checkbox(
    "Utiliser l'OCR (reconnaissance de texte)",
    value=True,
    help="Active la détection automatique de texte dans les images",
)

st.sidebar.markdown("---")
st.sidebar.info(
    "Information\n\n"
    "Cette application anonymise automatiquement :\n"
    "- Les dates (JJ/MM/AAAA)\n"
    "- Les numéros longs (6+ chiffres)\n"
    "- Les emails\n"
    "- Les numéros de téléphone\n"
    "- Les numéros de sécurité sociale\n"
    "- Les champs sélectionnés\n"
    "- Le texte dans les images (OCR)"
)

# ------------------------------------------------------------
# Zone principale : upload
# ------------------------------------------------------------
st.subheader("Charger le document médical")
uploaded_file = st.file_uploader(
    "Choisissez un fichier (PDF, Word, TXT ou Image)",
    type=["pdf", "docx", "doc", "txt", "png", "jpg", "jpeg", "gif", "bmp", "tiff"],
    help="Formats acceptés : PDF, DOCX, TXT, PNG, JPG, JPEG, GIF, BMP, TIFF",
)

if uploaded_file is not None:
    st.success(f"Fichier chargé : {uploaded_file.name}")

    file_extension = uploaded_file.name.split(".")[-1].lower()
    col1, col2 = None, None

    if file_extension in ["png", "jpg", "jpeg", "gif", "bmp", "tiff"]:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Image originale")
            st.image(uploaded_file, use_container_width=True)

    if st.button("Anonymiser le document", type="primary"):
        with st.spinner("Anonymisation en cours..."):
            try:
                file_bytes = uploaded_file.read()
                file_extension = uploaded_file.name.split(".")[-1].lower()

                if file_extension == "pdf":
                    anonymized_bytes, replacements = anonymize_pdf(
                        file_bytes, selected_labels
                    )
                    mime_type = "application/pdf"
                    output_extension = "pdf"

                elif file_extension in ["docx", "doc"]:
                    anonymized_bytes, replacements = anonymize_docx(
                        file_bytes, selected_labels
                    )
                    mime_type = (
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    )
                    output_extension = "docx"

                elif file_extension == "txt":
                    anonymized_bytes, replacements = anonymize_txt(
                        file_bytes, selected_labels
                    )
                    mime_type = "text/plain"
                    output_extension = "txt"

                elif file_extension in ["png", "jpg", "jpeg", "gif", "bmp", "tiff"]:
                    anonymized_bytes, replacements, img_format = anonymize_image(
                        file_bytes, selected_labels, use_ocr
                    )
                    mime_type = f"image/{img_format.lower()}"
                    output_extension = img_format.lower()

                else:
                    st.error("Type de fichier non pris en charge.")
                    st.stop()

                st.success("Anonymisation terminée !")

                if file_extension in ["png", "jpg", "jpeg", "gif", "bmp", "tiff"] and col2 is not None:
                    with col2:
                        st.subheader("Image anonymisée")
                        st.image(anonymized_bytes, use_container_width=True)

                col_stat1, col_stat2 = st.columns(2)
                with col_stat1:
                    st.metric("Éléments anonymisés", len(replacements))
                with col_stat2:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                if replacements:
                    st.subheader("Détails des anonymisations")
                    df_replacements = pd.DataFrame(
                        replacements,
                        columns=["Type", "Valeur originale", "Remplacement"],
                    )
                    st.dataframe(df_replacements, use_container_width=True)
                else:
                    st.info("Aucune donnée sensible détectée automatiquement.")

                st.subheader("Télécharger le document anonymisé")
                original_name = uploaded_file.name.rsplit(".", 1)[0]
                output_filename = f"{original_name}_anonymise_{timestamp}.{output_extension}"

                st.download_button(
                    label=f"Télécharger {output_filename}",
                    data=anonymized_bytes,
                    file_name=output_filename,
                    mime=mime_type,
                    type="primary",
                )

                st.warning(
                    "Attention : vérifiez toujours manuellement le document anonymisé "
                    "avant de le partager, pour vous assurer que toutes les données "
                    "sensibles ont été correctement supprimées."
                )

            except Exception as e:
                st.error(f"Erreur lors de l'anonymisation : {str(e)}")
                st.exception(e)

else:
    st.info(
        "Pour commencer :\n\n"
        "1. Sélectionnez les champs à anonymiser dans la barre latérale\n"
        "2. Téléchargez votre document médical (PDF, Word, TXT ou Image)\n"
        "3. Cliquez sur 'Anonymiser le document'\n"
        "4. Téléchargez le document anonymisé"
    )

with st.expander("Types de fichiers supportés"):
    st.markdown(
        """
    **Documents texte :**
    - PDF (avec masquage visuel des données)
    - Word (.docx)
    - Fichiers texte (.txt)

    **Images médicales :**
    - PNG
    - JPG / JPEG
    - GIF
    - BMP
    - TIFF

    Pour les images, l'OCR détecte automatiquement le texte et masque :
    - Les informations d'en-tête (nom, date, numéro)
    - Les dates et numéros dans l'image
    - Les zones de texte personnalisées
    """
    )

st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: gray;'>"
    "Application d'anonymisation de documents médicaux | "
    "Développée pour la protection des données patients | "
    "Support : PDF, Word, TXT, Images"
    "</div>",
    unsafe_allow_html=True,
)
